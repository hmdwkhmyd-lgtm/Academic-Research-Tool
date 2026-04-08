import wikipedia
import requests
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import io
import os

def create_final_academic_pro():
    # طلب اسم البحث بالإنجليزية لضمان دقة المصادر
    topic_in = input("Enter Research Title (English): ")
    print(f"🌐 جاري معالجة البحث الأكاديمي: {topic_in}...")
    
    doc = Document()
    trans = GoogleTranslator(source='en', target='ar')
    
    try:
        wikipedia.set_lang("en")
        search_res = wikipedia.search(topic_in)
        if not search_res:
            print("❌ لم يتم العثور على مراجع. حاول تبسيط العنوان.")
            return
            
        page = wikipedia.page(search_res[0], auto_suggest=False)
        
        # --- 1. صفحة غلاف أكاديمية (Professional Cover) ---
        doc.add_heading("Scientific Research Paper", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n" * 5)
        main_title = doc.add_heading(page.title, 0)
        main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n" * 8)
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Prepared by: خالد (IT Specialist)\nDate: 2026").font.size = Pt(16)
        doc.add_page_break()

        # --- 2. ملخص البحث (Abstract) ---
        doc.add_heading("ملخص البحث (Abstract)", level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        abstract = trans.translate(page.summary[:4000])
        p_abs = doc.add_paragraph(abstract)
        p_abs.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_abs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        doc.add_page_break()

        # --- 3. الفصول، الصور، والمحتوى ---
        sections = page.content.split('==')
        images = [i for i in page.images if i.endswith(('.jpg', '.png'))]
        img_idx = 0

        for block in sections:
            text = block.strip()
            if len(text) < 15: continue
            
            # إذا كان النص قصيراً فهو عنوان فصل
            if len(text) < 100:
                t_h = trans.translate(text)
                h = doc.add_heading(t_h, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run_h = h.add_run()
                run_h.font.color.rgb = RGBColor(0, 51, 102)
                
                # إضافة صورة توضيحية للفصل لزيادة الحجم والجمالية
                if img_idx < len(images):
                    try:
                        res = requests.get(images[img_idx], timeout=5)
                        doc.add_picture(io.BytesIO(res.content), width=Inches(4.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_idx += 1
                    except: pass
            else:
                # ترجمة وتنسيق الفقرات (تباعد أسطر مزدوج لزيادة الصفحات)
                t_p = trans.translate(text[:4500])
                p = doc.add_paragraph(t_p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.add_run().font.size = Pt(13)

        # --- 4. المراجع (References) ---
        doc.add_page_break()
        doc.add_heading("المراجع (References)", level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref = doc.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ref.add_run(f"• Wikipedia: {page.title} - Source: {page.url}\n")
        ref.add_run("• Automated Academic Retrieval System by Khaled (IT Student).")

        # --- 5. الحفظ في مجلد المخرجات ---
        output_path = "../Outputs/Academic_Research_Final.docx"
        doc.save(output_path)
        print(f"\n[✓] تم بنجاح! الملف موجود في: {output_path}")

    except Exception as e:
        print(f"❌ حدث خطأ: {e}")

if __name__ == "__main__":
    create_final_academic_pro()

