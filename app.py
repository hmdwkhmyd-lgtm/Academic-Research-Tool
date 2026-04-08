import streamlit as st
import wikipedia
import requests
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import io

# --- إعدادات الصفحة ---
st.set_page_config(page_title="صانع البحوث الأكاديمي", page_icon="🎓")

# --- الواجهة الجانبية (Sidebar) ---
with st.sidebar:
    st.title("👨‍💻 المبرمج خالد")
    st.info("طالب IT متخصص في أتمتة الحلول البرمجية")
    st.markdown("---")
    st.write("هذه الأداة تقوم بسحب البيانات من المصادر الدولية وتنسيقها أكاديمياً.")

# --- الواجهة الرئيسية ---
st.title("🚀 نظام إنشاء البحوث الأكاديمية الدولي")
st.subheader("حول فكرتك إلى بحث متكامل بضغطة زر")

topic_in = st.text_input("أدخل عنوان البحث بالإنجليزية (مثال: Internet of Things):")

if st.button("بدء تصميم البحث ✨"):
    if topic_in:
        with st.spinner('جاري سحب البيانات، الترجمة، والتنسيق...'):
            try:
                doc = Document()
                trans = GoogleTranslator(source='en', target='ar')
                wikipedia.set_lang("en")
                
                search_res = wikipedia.search(topic_in)
                if not search_res:
                    st.error("لم يتم العثور على مراجع لهذا الموضوع.")
                else:
                    page = wikipedia.page(search_res[0], auto_suggest=False)
                    
                    # --- الغلاف ---
                    doc.add_heading("Academic Research Paper", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("\n" * 5)
                    title = doc.add_heading(page.title, 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("\n" * 8)
                    doc.add_paragraph(f"Prepared by: Khaled (IT)\nDate: 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_page_break()

                    # --- المحتوى ---
                    sections = page.content.split('==')
                    images = [img for img in page.images if img

